from navec import Navec
from slovnet import NER
from slovnet.span import Span
import docx
import re
from functools import total_ordering
from slovnet.markup import SpanMarkup
import textract
from model.model import fit_time

class DOC:
    dates=None
    def __init__(self, file_name=None, text=None, paragraphs=None, per=None, processed=None, tags_deletion=False):
        if (processed == True):
            self.text = text
            self.paragraphs = paragraphs
            self.per = per
        elif (file_name is not None):
            if (not tags_deletion):
                other = DOC_from_text(docx_to_text(file_name))
            else:
                other = DOC_from_text(docx_to_text(file_name), date_tags_deletion=True)
            self.text = other.text
            self.paragraphs = other.paragraphs
            self.per = other.per
    def dates_extraction(self):
        self.dates = fit_time(self.text)




def docx_to_text(file_name):
    return textract.process(file_name).decode('utf8')

def DOC_from_text(text, date_tags_deletion=False):
    doc_ = "" #text of the document
    spans_per = []
    spans_paragraphs = []


    reg_exp = r"(<персона.*?>)(.*?)(</персона>)"
    reg_exp_date = r"<дата.*?>"

    @total_ordering
    class event:
        __attributes__ = ['span', 'type', 'position']

        def __init__(self, span, type, position=None):
            self.span = span
            self.type = type
            self.position = position
        def __str__(self):
            s = str(self.span) + " " + self.type
            if (self.position is not None):
                s += " " + self.position
            return s
        def __eq__(self, other):
            return ((self.span, self.type, self.position) == (other.span, other.type, other.position))
        def __ne__(self, other):
            return not (self == other)
        def __lt__(self, other):
            return (self.span[0] < other.span[0])


    events = []

    m = re.finditer(reg_exp, text)

    for s in m:
        if (s.span(2)[0] < s.span(2)[1]):    # markup leaves wishing for better
            events.append(event(s.span(1), 'PER', 'start'))
            events.append(event(s.span(2), 'PER', 'middle'))
            events.append(event(s.span(3), 'PER', 'end'))

    date = re.finditer(reg_exp_date, text)

    for a in date:
        # print(text[a.span()[0]:a.span()[1]])
        events.append(event(a.span(), 'Date'))

    events.sort()

    i = 0
    j = 0

    prev = 0
    is_first=True

    def strip(text, start, end):
        while (text[start] in ['\n', ' '] and start + 1 < end):
            start += 1
        while (text[end - 1] in ['\n', ' '] and start + 1 < end):
            end -= 1

        return (start, end)

    while(i < len(text)):
        if (j < len(events) and events[j].span[0] == i):
            if (events[j].type == 'PER'):
                if (events[j].position == 'start' or events[j].position == 'end'):
                    i = events[j].span[1]
                elif (events[j].position == 'middle'):
                    start = events[j].span[0]
                    end = events[j].span[1]

                    start, end = strip(text, start, end)

                    spans_per.append(Span(len(doc_), len(doc_) + end - start, 'PER'))
                    doc_ += text[start : end]
                    i = events[j].span[1]
            elif (events[j].type == 'Date'):
                # print('Date')
                # strip
                end = len(doc_)
                # print(prev, end)
                # print(doc_[prev:end])
                if (not is_first):
                    prev, end = strip(doc_, prev, end)
                    spans_paragraphs.append(Span(prev, end, 'Paragraph'))
                if (not date_tags_deletion):
                    doc_ += text[events[j].span[0] : events[j].span[1]]
                i = events[j].span[1]
                is_first = False
                prev = len(doc_)
            j += 1
        else:
            doc_ += text[i]
            i += 1

    end = len(doc_)

    prev, end = strip(doc_, prev, end)

    spans_paragraphs.append(Span(prev, end, 'Paragraph'))

    for p in spans_paragraphs:
        if (p.start >= p.stop):
            raise NameError("Irregular span. Start after stop.")
    for p in spans_per:
        if (p.start >= p.stop):
            raise NameError("Irregular span. PER")

    return DOC(text=doc_, paragraphs=spans_paragraphs, per=spans_per, processed=True)